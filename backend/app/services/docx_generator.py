"""Generate the technical specification Word document from a QuotationResponse + request."""
from io import BytesIO
from datetime import date

from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from models.db import Breaker
from models.schemas import QuotationResponse, QuotationRequest
from data.components import PANEL_DEFAULTS
from data.fixed_lookups import (
    BUSBAR_THRESHOLD_AMPERAGE,
    cable_size_for_amperage,
    outgoing_busbar_for_amperage,
)


HEADER_BLUE = RGBColor(0x1F, 0x4E, 0x78)


def _set_cell_background(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = HEADER_BLUE
        run.font.name = "Arial"
    return h


def _set_default_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)


def _add_kv_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(4.0)

    for label, value in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value)
        for run in row[0].paragraphs[0].runs:
            run.bold = True
        _set_cell_background(row[0], "D9E2F3")
    return table


def _header_row(table, headers: list[str]):
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
        _set_cell_background(hdr[i], "1F4E78")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _breaker_table_rows(doc: Document, db: Session, specs):
    """Build a table of breaker specs with resolved cable size / busbar."""
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    _header_row(table, ["Quantity", "OEM", "Series", "Amperage (A)", "Rating (kA)", "Cable Size / Busbar"])

    for spec in specs:
        breaker = db.query(Breaker).filter(Breaker.id == spec.breaker_id).first()
        if not breaker:
            continue

        if breaker.amperage < BUSBAR_THRESHOLD_AMPERAGE:
            supply = cable_size_for_amperage(breaker.amperage) or "-"
        else:
            result = outgoing_busbar_for_amperage(breaker.amperage)
            supply = result[1] if result else "-"

        row = table.add_row().cells
        row[0].text = str(spec.quantity)
        row[1].text = breaker.oem
        row[2].text = breaker.series
        row[3].text = str(breaker.amperage)
        row[4].text = str(breaker.rating_kA)
        row[5].text = supply

    return table


def generate_spec_docx(quotation: QuotationResponse, req: QuotationRequest, db: Session) -> bytes:
    doc = Document()
    _set_default_font(doc)

    # Title
    title = doc.add_heading("AC Combiner Panel", level=0)
    for run in title.runs:
        run.font.color.rgb = HEADER_BLUE
        run.font.name = "Arial"

    sub = doc.add_paragraph("Technical Specification")
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph(f"Customer: {quotation.customer_name}")
    doc.add_paragraph(f"Project: {quotation.project_name}")
    doc.add_paragraph(f"OEM: {quotation.oem}")
    doc.add_paragraph(f"Date: {date.today().isoformat()}")
    doc.add_paragraph()

    # --- Panel configuration summary ---
    _add_heading(doc, "1. Panel Configuration Summary", level=1)
    summary_rows = [
        ("Customer", quotation.customer_name),
        ("OEM (Breaker Supplier)", quotation.oem),
        ("Number of Incomer Circuits", str(sum(s.quantity for s in req.incomers))),
        ("Number of Outgoing Circuits", str(sum(s.quantity for s in req.outgoings))),
        ("Enclosure Rating", PANEL_DEFAULTS["enclosure_rating"]),
        ("Enclosure Material", PANEL_DEFAULTS["enclosure_material"]),
    ]
    _add_kv_table(doc, summary_rows)
    doc.add_paragraph()

    # --- Electrical specifications ---
    _add_heading(doc, "2. Electrical Specifications", level=1)
    elec_rows = [
        ("System Voltage", PANEL_DEFAULTS["system_voltage"]),
        ("Control Voltage", PANEL_DEFAULTS["control_voltage"]),
        ("Busbar Material", PANEL_DEFAULTS["busbar_material"]),
        ("Earth Fault Protection", PANEL_DEFAULTS["earth_fault_protection"]),
        ("Applicable Standards", PANEL_DEFAULTS["standard_compliance"]),
    ]
    _add_kv_table(doc, elec_rows)
    doc.add_paragraph()

    # --- Incomer breaker specifications ---
    _add_heading(doc, "3. Incomer Breaker Specifications", level=1)
    _breaker_table_rows(doc, db, req.incomers)
    doc.add_paragraph()

    # --- Outgoing breaker specifications ---
    _add_heading(doc, "4. Outgoing Breaker Specifications", level=1)
    _breaker_table_rows(doc, db, req.outgoings)
    doc.add_paragraph()

    # --- Main busbar ---
    main_busbar_lines = [l for l in quotation.bom if l.category == "busbar" and "main busbar" in l.description.lower()]
    if main_busbar_lines:
        _add_heading(doc, "5. Main Busbar", level=1)
        for line in main_busbar_lines:
            doc.add_paragraph(line.description, style="List Bullet")
        doc.add_paragraph()
        section_num = 6
    else:
        section_num = 5

    # --- Component / accessory details ---
    _add_heading(doc, f"{section_num}. Component & Accessory Details", level=1)
    accessory_lines = [l for l in quotation.bom if l.category == "accessory"]
    if accessory_lines:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        _header_row(table, ["Component", "Part Number", "Quantity"])

        for line in accessory_lines:
            row = table.add_row().cells
            row[0].text = line.description
            row[1].text = line.part_number
            row[2].text = str(line.quantity)
    else:
        doc.add_paragraph("No additional accessories included in this configuration.")

    doc.add_paragraph()
    section_num += 1

    # --- Notes ---
    if quotation.warnings:
        _add_heading(doc, f"{section_num}. Notes & Exceptions", level=1)
        for w in quotation.warnings:
            doc.add_paragraph(w, style="List Bullet")

    # --- Footer note ---
    doc.add_paragraph()
    footer = doc.add_paragraph(
        "This document is system-generated based on the configuration selected at time of "
        "quotation and is intended for budgetary / proposal purposes. Final design is subject "
        "to engineering review and site survey."
    )
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.italic = True
    footer.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
